# -*- coding: utf-8 -*-
"""Render a Markdown doc (headings, bullets, numbered lists, code blocks,
inline **bold**/`code`, and pipe tables) to a readable .docx.

Usage: md2docx.py <src.md> <dst.docx> [Document Title]
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC, DST = sys.argv[1], sys.argv[2]
TITLE = sys.argv[3] if len(sys.argv) > 3 else None

MONO = "Consolas"
ACCENT = RGBColor(0xB4, 0x4A, 0x1E)
GREY = RGBColor(0x80, 0x80, 0x80)
CODECLR = RGBColor(0x6B, 0x3A, 0x1F)

doc = Document()
n = doc.styles["Normal"]
n.font.name = "Calibri"
n.font.size = Pt(11)


def add_runs(par, text):
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|`([^`]+)`", text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            par.add_run(m.group(1)).bold = True
        else:
            r = par.add_run(m.group(2))
            r.font.name = MONO
            r.font.size = Pt(10)
            r.font.color.rgb = CODECLR
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]+$", line)) and "-" in line


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


if TITLE:
    h = doc.add_heading(TITLE, level=0)
    for r in h.runs:
        r.font.color.rgb = ACCENT

with open(SRC, encoding="utf-8") as f:
    lines = f.read().splitlines()

i = 0
in_code = False
while i < len(lines):
    line = lines[i]

    if line.startswith("```"):
        in_code = not in_code
        i += 1
        continue
    if in_code:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(line)
        r.font.name = MONO
        r.font.size = Pt(9)
        i += 1
        continue

    # Table: header line, separator, then body rows
    if ("|" in line and i + 1 < len(lines) and is_table_sep(lines[i + 1])
            and not line.lstrip().startswith(">")):
        header = split_row(line)
        i += 2
        body = []
        while i < len(lines) and "|" in lines[i] and lines[i].strip():
            body.append(split_row(lines[i]))
            i += 1
        ncol = len(header)
        t = doc.add_table(rows=1, cols=ncol)
        t.style = "Light Grid Accent 2"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for j, cell in enumerate(header):
            c = t.rows[0].cells[j]
            c.paragraphs[0].text = ""
            add_runs(c.paragraphs[0], cell)
            for r in c.paragraphs[0].runs:
                r.bold = True
        for row in body:
            cells = t.add_row().cells
            for j in range(ncol):
                val = row[j] if j < len(row) else ""
                cells[j].paragraphs[0].text = ""
                add_runs(cells[j].paragraphs[0], val)
                for r in cells[j].paragraphs[0].runs:
                    r.font.size = Pt(9)
        doc.add_paragraph()
        continue

    if line.startswith("# "):
        h = doc.add_heading(line[2:], level=0)
        for r in h.runs:
            r.font.color.rgb = ACCENT
        i += 1
        continue
    if line.startswith("## "):
        doc.add_heading(line[3:], level=1)
        i += 1
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:], level=2)
        i += 1
        continue
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(re.sub(r"`([^`]+)`", r"\1", line[2:]))
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        i += 1
        continue
    if line.strip() == ">":
        i += 1
        continue
    if re.match(r"^\s*- ", line):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, re.sub(r"^\s*- ", "", line))
        i += 1
        continue
    m = re.match(r"^\s*(\d+)\. (.*)$", line)
    if m:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(m.group(1) + ". ").bold = True
        add_runs(p, m.group(2))
        i += 1
        continue
    if line.strip() == "":
        i += 1
        continue
    p = doc.add_paragraph()
    add_runs(p, line)
    i += 1

doc.save(DST)
print("saved:", DST)
